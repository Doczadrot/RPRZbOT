# neuro_rag.py
# Модуль для реализации функционала "Нейропомощника РПРЗ" с использованием RAG (Retrieval-Augmented Generation)
# Синхронная реализация для интеграции с pyTelegramBotAPI.

# --- Импорт необходимых библиотек --- ✨
import os # Для работы с файловой системой (папки, файлы)
import requests # Для выполнения веб-запросов (общение с API нейросети)
import json # Для работы с данными в формате JSON (часто используется в API)
# import openai # Официальная библиотека для работы с OpenAI API (будет использоваться через Langchain)
# import pickle # Может понадобиться для сохранения/загрузки FAISS индекса
from dotenv import load_dotenv # Для загрузки переменных окружения из файла .env

# Импорты для чтения разных форматов документов
from docx import Document # Для работы с .docx файлами
from PyPDF2 import PdfReader # Для работы с .pdf файлами (альтернатива: pdfminer.six)

# Импорты из Langchain и других библиотек для RAG
# Эти библиотеки помогут нам обрабатывать текст, создавать векторы и искать по ним.
from langchain_text_splitters import RecursiveCharacterTextSplitter # Для разделения текста на чанки

# ✅ ИМПОРТЫ ДЛЯ ЭМБЕДДИНГОВ И LLM (теперь используем OpenAI-совместимые классы, указывая URL DeepSeek)
# from langchain_community.embeddings import DeepseekEmbeddings # <-- Этот импорт больше не нужен
from langchain_openai import OpenAIEmbeddings # ✅ Используем OpenAIEmbeddings для эмбеддингов DeepSeek
from langchain_community.vectorstores import FAISS # Для локальной векторной базы данных FAISS

# Импорт для работы с Document объектами (часть Langchain)
from langchain_core.documents import Document

# ✅ Импорт для клиента DeepSeek LLM через OpenAI-совместимый клиент Langchain
from langchain_openai import ChatOpenAI # ✅ Используем ChatOpenAI для чат-модели DeepSeek
# from langchain_community.chat_models import ChatDeepseek # <-- Этот импорт больше не нужен


<<<<<<< HEAD
# --- Настройка и Конфигурация --- ⚙️
load_dotenv() # ✅ Загружаем переменные окружения из файла .env. Убедитесь, что файл .env существует!

# --- Ключ API и URL для DeepSeek (используем переменные из .env) --- 🔑
DEEPSEEK_API_KEY = os.getenv('DEEPSEEK_API_KEY')
DEEPSEEK_API_BASE = os.getenv('DEEPSEEK_API_BASE') # ✅ Читаем базовый URL из .env

if not DEEPSEEK_API_KEY:
    print("Ошибка: DEEPSEEK_API_KEY не найден в файле .env!")
    # В реальном проекте здесь можно вызвать sys.exit(1) или поднять исключение
elif not DEEPSEEK_API_BASE:
     print("Ошибка: DEEPSEEK_API_BASE не найден в файле .env!")
else:
    print("[DEBUG] Ключ и базовый URL DeepSeek API загружены.") # Лог подтверждения

# --- Названия моделей DeepSeek (используем переменные из .env) --- 🤖🧠
LLM_MODEL_NAME = os.getenv('DEEPSEEK_LLM_MODEL_NAME', 'deepseek-chat') # ✅ Модель для ответов бота DeepSeek
EMBEDDING_MODEL_NAME = os.getenv('DEEPSEEK_EMBEDDING_MODEL_NAME', 'deepseek-text-embedding') # ✅ Модель для эмбеддингов DeepSeek

# Путь к папке, где администратор будет хранить загруженные вручную документы. 📂
DOCUMENTS_DIR = os.getenv('DOCUMENTS_DIR', './documents') # ✅ Путь лучше брать из .env

# Путь к папке, где будет сохраняться индекс базы знаний FAISS (векторы документов). 📊
INDEX_PERSIST_DIR = os.getenv('INDEX_PERSIST_DIR', './vector_index') # ✅ Путь лучше брать из .env


# Список ID пользователей Telegram, которые являются администраторами бота. 👑
# Рекомендуется хранить в .env в виде строки, например: ADMIN_USER_IDS='123456789,987654321'
ADMIN_USER_IDS_STR = os.getenv('ADMIN_USER_IDS', '')
ADMIN_USER_IDS = [int(uid.strip()) for uid in ADMIN_USER_IDS_STR.split(',') if uid.strip().isdigit()] if ADMIN_USER_IDS_STR else []
if not ADMIN_USER_IDS_STR: # Проверяем исходную строку из .env
    print("Внимание: Переменная ADMIN_USER_IDS не найдена в файле .env или пуста!")
elif not ADMIN_USER_IDS: # Проверяем, удалось ли получить хотя бы один корректный ID
    print(f"Внимание: Не удалось извлечь корректные числовые ID администраторов из строки: {ADMIN_USER_IDS_STR}")
else:
    print(f"[DEBUG] Загружены ID администраторов: {ADMIN_USER_IDS}") # Лог загруженных ID


# --- Создаем папки при запуске скрипта --- ✨
os.makedirs(DOCUMENTS_DIR, exist_ok=True) # ✅ Создание папки для документов.
os.makedirs(INDEX_PERSIST_DIR, exist_ok=True) # ✅ Создание папки для индекса базы знаний.
print(f"[DEBUG] Проверены папки: {DOCUMENTS_DIR}, {INDEX_PERSIST_DIR}") # Лог подтверждения

# --- Инициализация компонентов RAG --- 🧠📊
# Инициализируем модель для создания эмбеддингов, используя OpenAIEmbeddings, но направляя их в DeepSeek
# ✅ ИСПРАВЛЕНО: Использование OpenAIEmbeddings с указанием базового URL DeepSeek
embeddings_model = OpenAIEmbeddings(
    openai_api_key=DEEPSEEK_API_KEY, # ✅ Используем ключ DeepSeek
    model=EMBEDDING_MODEL_NAME,     # ✅ Используем название модели эмбеддингов DeepSeek
    base_url=DEEPSEEK_API_BASE      # ✅ Указываем базовый URL DeepSeek API
)
print(f"[DEBUG] Инициализирована модель эмбеддингов (OpenAI-совместимая) DeepSeek: {EMBEDDING_MODEL_NAME}") # Лог подтверждения

# ✅ Инициализация клиента DeepSeek LLM (понадобится позже для Шага 8)
# Используем ChatOpenAI, указывая базовый URL DeepSeek
llm_client = ChatOpenAI(
    openai_api_key=DEEPSEEK_API_KEY, # ✅ Используем ключ DeepSeek
    model=LLM_MODEL_NAME,          # ✅ Используем название модели чата DeepSeek
    base_url=DEEPSEEK_API_BASE,    # ✅ Указываем базовый URL DeepSeek API
    temperature=0.0
)
print(f"[DEBUG] Инициализирована чат-модель (OpenAI-совместимая) DeepSeek: {LLM_MODEL_NAME}") # Лог подтверждения


# ШАГ 2: Загрузка Текста из Файлов ✨
def load_document_text(file_path: str) -> str | None:
    """
    Читает текст из файла в зависимости от его расширения.
    Поддерживает .pdf, .docx, .txt.
    Возвращает текст документа или None, если формат не поддерживается или произошла ошибка.
    """
    text = ""
    try:
        # Получаем расширение файла в нижнем регистре (например, '.pdf')
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower() # Приводим к нижнему регистру для надежности

        if file_extension == '.pdf':
            # --- Чтение PDF ---
            with open(file_path, 'rb') as f: # Открываем файл в бинарном режиме для чтения
                reader = PdfReader(f) # Создаем объект PdfReader
                for page_num in range(len(reader.pages)): # Проходим по всем страницам
                    page = reader.pages[page_num] # Получаем одну страницу
                    text += page.extract_text() or "" # Извлекаем текст со страницы и добавляем к общему тексту
            print(f"[DEBUG] Успешно прочитан PDF: {os.path.basename(file_path)}")
            return text
        elif file_extension == '.docx':
            # --- Чтение DOCX ---
            doc = Document(file_path) # Открываем .docx файл
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs]) # Собираем текст из всех параграфов
            print(f"[DEBUG] Успешно прочитан DOCX: {os.path.basename(file_path)}")
            return text
        elif file_extension == '.txt':
            # --- Чтение TXT ---
            with open(file_path, 'r', encoding='utf-8') as f: # Открываем текстовый файл с кодировкой UTF-8
                text = f.read() # Читаем весь текст
            print(f"[DEBUG] Успешно прочитан TXT: {os.path.basename(file_path)}")
            return text
        else:
            # Если формат файла не поддерживается
            print(f"[WARNING] Неподдерживаемый формат файла: {file_path}")
            return None # Возвращаем None для неподдерживаемых форматов
    except Exception as e:
        # Если произошла любая другая ошибка при чтении файла
        print(f"[ERROR] Ошибка при чтении файла {file_path}: {e}")
        return None

# Функция для загрузки векторной базы (использует embeddings_model)
# ✅ Эта функция нужна для Шага 6 и Шага 9
def load_vector_store(vector_store_dir: str, embeddings):
    """Загружает существующий векторный индекс FAISS с диска."""
    print(f"[DEBUG] Попытка загрузки векторного хранилища из {vector_store_dir}...")
    # Проверяем наличие основного файла индекса (index.faiss)
    index_file = os.path.join(vector_store_dir, "index.faiss")
    # ✅ Исправлено: Проверяем существование папки перед попыткой загрузки
    if not os.path.exists(vector_store_dir) or not os.path.exists(index_file):
        print(f"[DEBUG] Файл индекса FAISS не найден по пути {index_file}. Не удалось загрузить vector store.")
        return None
    try:
        # Загружаем индекс с диска, передавая модель эмбеддингов (теперь DeepseekEmbeddings через OpenAIEmbeddings)
        vector_store = FAISS.load_local(
            folder_path=vector_store_dir,
            embeddings=embeddings, # ✅ Передаем модель эмбеддингов (которая теперь DeepseekEmbeddings через OpenAIEmbeddings)
            allow_dangerous_deserialization=True # Нужно для загрузки из файла
        )
        print(f"[DEBUG] Vector store успешно загружен из: {vector_store_dir}")
        return vector_store
    except Exception as e:
        print(f"[ERROR] Ошибка при загрузке vector store: {e}")
        return None


# ШАГ 6 (упрощенный): Обновление Базы Знаний из папки (включает ШАГ 3, 4, 5)
def update_knowledge_base():
    """
    Сканирует папку DOCUMENTS_DIR, читает текст из всех поддерживаемых файлов,
    делит его на чанки, создает эмбеддинги и обновляет (или создает)
    векторную базу FAISS в папке INDEX_PERSIST_DIR.
    Возвращает True в случае успеха, False в случае ошибки.
    """
    print(f"[DEBUG] Начинаем обновление базы знаний из папки: {DOCUMENTS_DIR}") # Лог: начинаем процесс

    # Получаем список всех элементов (файлов и папок) в директории документов
    try:
        entries = os.listdir(DOCUMENTS_DIR) # Получаем список имен файлов и папок в папке
        print(f"[DEBUG] Найдено элементов в папке: {len(entries)}") # Лог: сколько всего нашли
    except FileNotFoundError:
        print(f"[ERROR] Папка документов не найдена: {DOCUMENTS_DIR}") # Лог ошибки: папка не найдена
        return False # Возвращаем False, если папка не найдена

    # Список для хранения полных путей только к файлам, которые будем обрабатывать
    file_paths = []

    # Проходимся по всем найденным элементам в папке
    for entry_name in entries: # entry_name - это просто имя файла или папки
        full_path = os.path.join(DOCUMENTS_DIR, entry_name) # Получаем полный путь к элементу

        if os.path.isfile(full_path): # Если это файл...
            file_paths.append(full_path) # ...добавляем его полный путь в наш список файлов

    print(f"[DEBUG] Найдено поддерживаемых файлов для обработки: {len(file_paths)}") # Лог: сколько файлов будем обрабатывать

    if not file_paths:
        print("[DEBUG] Нет файлов для обработки. Обновление базы знаний не требуется.")
        # Попробуем загрузить существующий индекс даже если нет новых файлов
        # Это нужно, если бот перезапустился, а индекс уже был создан
        print("[DEBUG] Пробуем загрузить существующий индекс FAISS.")
        vector_store = load_vector_store(INDEX_PERSIST_DIR, embeddings_model) # <-- Используем load_vector_store
        if vector_store:
             print("[DEBUG] Индекс FAISS успешно загружен (нет новых файлов для добавления).")
             # ✅ Теперь функция возвращает сам vector_store объект при успешной загрузке/создании
             return vector_store # Успех, хотя новых файлов не было
        else:
             print("[DEBUG] Существующий индекс не найден или не может быть загружен. База знаний пуста.")
             return False # Нет файлов для индексирования и нет старого индекса


    # --- ШАГ 6.2 (упрощенный): Читаем текст из найденных файлов --- ✨
    # Список для хранения успешно загруженных документов (кортеж: путь к файлу, текст)
    loaded_documents_text = [] # Создаем пустой список для результатов

    # Проходимся по каждому найденному пути к файлу и читаем его текст
    for file_path in file_paths: # Перебираем (путь, текст) из списка прочитанных документов
        print(f"[DEBUG] Читаем файл: {file_path}") # Лог: какой файл сейчас читаем

        document_text = load_document_text(file_path) # Вызываем функцию для чтения текста

        if document_text: # Если load_document_text вернула текст (не None и не пустую строку)...
            loaded_documents_text.append((file_path, document_text)) # ...добавляем результат в список
            print(f"[DEBUG] Текст из {file_path} успешно прочитан.") # Лог: успешно прочитали
        else:
            # Лог об ошибке или неподдерживаемом формате уже есть внутри load_document_text
            print(f"[DEBUG] Пропуск файла {file_path}: не удалось прочитать текст или формат не поддерживается.") # Лог: не удалось прочитать

    print(f"[DEBUG] Всего успешно прочитано файлов: {len(loaded_documents_text)}") # Лог: сколько файлов удалось прочитать

    if not loaded_documents_text:
        print("[DEBUG] Нет прочитанного текста из файлов. Индексирование невозможно.")
        # Если нет текста для индексирования, все равно пробуем загрузить старый индекс
        print("[DEBUG] Пробуем загрузить существующий индекс FAISS, так как нет нового текста.")
        vector_store = load_vector_store(INDEX_PERSIST_DIR, embeddings_model) # <-- Используем load_vector_store
        if vector_store:
             print("[DEBUG] Индекс FAISS успешно загружен (нет нового текста для добавления).")
             # ✅ Функция возвращает сам vector_store объект
             return vector_store # Успех, хотя индексация не проводилась
        else:
             print("[DEBUG] Существующий индекс не найден или не может быть загружен. База знаний пуста.")
             return False # Нет текста и нет старого индекса


    # --- ШАГ 3 + ШАГ 4 + ШАГ 5: Разбиение на чанки, Векторизация и Индексирование --- ✨
    print("[DEBUG] Начинаем разделение на чанки, векторизацию и индексирование...")

    # Инициализируем делитель текста.
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200) # Создаем делитель

    # Список для хранения всех полученных чанков в виде объектов Document
    all_document_chunks = [] # Создаем пустой список для Document объектов

    # Проходимся по каждому прочитанному документу (путь, текст)
    for file_path, document_text in loaded_documents_text:
        print(f"[DEBUG] Обрабатываем документ для индексирования: {file_path}")

        # Разбиваем текст текущего документа на чанки (строки)
        chunks_text = text_splitter.split_text(document_text) # Разбиваем текст на чанки (строки)

        # Превращаем строки-чанки в объекты Document с метаданными
        # Метаданные помогают отслеживать, откуда взялся чанк.
        # ИСПРАВЛЕНО: убран keyword page_content, как мы делали ранее.
        document_chunks = [
            Document(chunk, metadata={"source": os.path.basename(file_path)}) # ✅ Исправлено здесь
            for chunk in chunks_text
        ]
        all_document_chunks.extend(document_chunks) # Добавляем эти Document объекты в общий список


    print(f"[DEBUG] Всего получено чанков (объектов Document): {len(all_document_chunks)}")

    if not all_document_chunks:
        print("[DEBUG] Нет чанков для индексирования.")
        return False # Ошибка, если нет чанков

    # --- Индексирование в FAISS ---
    vector_store = None
    # Проверяем, существует ли уже индекс FAISS на диске
    # Проверяем наличие основного файла индекса (index.faiss)
    faiss_index_file = os.path.join(INDEX_PERSIST_DIR, 'index.faiss')
    # ✅ Добавлена проверка существования папки перед попыткой загрузки индекса
    if os.path.exists(INDEX_PERSIST_DIR) and os.path.exists(faiss_index_file):
        print(f"[DEBUG] Загружаем существующий индекс FAISS из: {INDEX_PERSIST_DIR}")
        # ✅ ИСПРАВЛЕНО: Передаем embeddings_model в load_local
        vector_store = load_vector_store(INDEX_PERSIST_DIR, embeddings_model) # ✅ ИСПРАВЛЕНО: Используем load_vector_store

        if vector_store:
            # Если индекс загружен, добавляем в него новые чанки
            print(f"[DEBUG] Добавляем {len(all_document_chunks)} новых чанков в существующий индекс...")
            try:
                # Для FAISS.add_documents embeddings НЕ передаются.
                vector_store.add_documents(all_document_chunks) # ✅ Добавляем Document объекты
                print("[DEBUG] Новые чанки успешно добавлены в индекс.")
            except Exception as e:
                print(f"[ERROR] Ошибка при добавлении чанков в индекс: {e}")
                print("[ERROR] Индекс не был обновлен.")
                # В случае ошибки добавления, текущий индекс (если был загружен) остается неизменным
                return False # Ошибка при добавлении
        else:
             # Если не удалось загрузить существующий индекс, нужно понять, хотим ли мы его перезаписать.
             # Пока просто логируем ошибку загрузки выше и возвращаем False из load_vector_store
             # Если load_vector_store вернул None, значит, не удалось загрузить старый индекс.
             print("[ERROR] Не удалось загрузить существующий индекс FAISS.")
             # Переходим к созданию нового индекса
             vector_store = None # Убеждаемся, что vector_store = None

    # Если индекс не существует или не удалось загрузить, создаем новый
    if vector_store is None: # ✅ Проверяем, что vector_store = None после попытки загрузки или изначально
        print(f"[DEBUG] Создаем новый индекс FAISS из {len(all_document_chunks)} чанков...")
        if not all_document_chunks:
             print("[DEBUG] Нет чанков для создания нового индекса.")
             return False # Нельзя создать индекс без чанков

        try:
            # Создаем новый индекс из всех Document объектов и их эмбеддингов DeepSeek
            # ✅ Передаем embeddings_model (который теперь OpenAIEmbeddings с базовым URL DeepSeek)
            vector_store = FAISS.from_documents(
                documents=all_document_chunks, # Передаем Document объекты
                embedding=embeddings_model # ✅ Передаем модель эмбеддингов
            )
            print("[DEBUG] Новый индекс FAISS успешно создан.")
        except Exception as e:
            print(f"[ERROR] Ошибка при создании нового индекса FAISS: {e}")
            # ✅ Лог ошибки API
            print(f"[ERROR] Детали ошибки API: {e}")
            print("[ERROR] База знаний не была создана.")
            return False # Ошибка при создании


    # --- Сохранение индекса на диск ---
    # Убедимся, что папка для сохранения существует (хотя мы создаем ее при запуске скрипта)
    os.makedirs(INDEX_PERSIST_DIR, exist_ok=True)
    print(f"[DEBUG] Сохраняем индекс FAISS в: {INDEX_PERSIST_DIR}")
    try:
        # Сохраняем индекс на диск
        vector_store.save_local(INDEX_PERSIST_DIR) # ✅ Сохраняем индекс на диск
        print("[DEBUG] Индекс FAISS успешно сохранен.")
        # ✅ Функция теперь возвращает сам vector_store объект при успешной индексации
        return vector_store # ✅ Успех: база знаний обновлена/создана и сохранена! Возвращаем объект индекса.
    except Exception as e:
        print(f"[ERROR] Ошибка при сохранении индекса FAISS: {e}")
        print("[ERROR] Индекс не был сохранен.")
        return False # Ошибка при сохранении


# ✅ Добавляем функцию для загрузки вексторного хранилища отдельно
# Эта функция уже была внутри update_knowledge_base и тестового блока, но лучше определить ее отдельно
def load_vector_store(vector_store_dir: str, embeddings):
    """Загружает существующий векторный индекс FAISS с диска."""
    print(f"[DEBUG] Попытка загрузки векторного хранилища из {vector_store_dir}...")
    # Проверяем наличие основного файла индекса (index.faiss)
    index_file = os.path.join(vector_store_dir, "index.faiss")
    # ✅ Исправлено: Проверяем существование папки перед попыткой загрузки
    if not os.path.exists(vector_store_dir) or not os.path.exists(index_file):
        print(f"[DEBUG] Файл индекса FAISS не найден по пути {index_file}. Не удалось загрузить vector store.")
        return None
    try:
        # Загружаем индекс с диска, передавая модель эмбеддингов (теперь DeepseekEmbeddings через OpenAIEmbeddings)
        vector_store = FAISS.load_local(
            folder_path=vector_store_dir,
            embeddings=embeddings, # ✅ Передаем модель эмбеддингов
            allow_dangerous_deserialization=True # Нужно для загрузки из файла
        )
        print(f"[DEBUG] Vector store успешно загружен из: {vector_store_dir}")
        return vector_store
    except Exception as e:
        print(f"[ERROR] Ошибка при загрузке vector store: {e}")
        return None


# --- Здесь будут другие функции (query_rag и т.д.) ---

# Временно вызываем функцию обновления базы при запуске скрипта для проверки
if __name__ == "__main__":
    print("\n--- Запуск скрипта neuro_rag.py ---")
    # ✅ Теперь update_knowledge_base возвращает объект vector_store или False
    vector_store_object = update_knowledge_base()
    if vector_store_object: # Если функция вернула объект vector_store (успех)
        print("\n--- Обновление базы знаний завершено успешно! ---")
        print(f"База знаний FAISS создана/обновлена в папке: {INDEX_PERSIST_DIR}")
        print("Теперь можно использовать объект vector_store_object для поиска.")
        # ✅ vector_store_object уже содержит загруженный/созданный индекс

        # Добавляем тестовый поиск здесь
        print("\n--- Тестовый поиск ---")
        query = "Какой защитный слой должен быть у фундаментов?" # Ваш тестовый запрос
        print(f"Запрос: {query}")
        # Внимание: similarity_search может работать чуть иначе,
        # в зависимости от того, как создан индекс. Но similarity_search должна работать.
        try:
            # Используем vector_store_object, который мы получили из update_knowledge_base
            relevant_docs = vector_store_object.similarity_search(query, k=3) # Ищем 3 самых похожих чанка
            print("Найденные релевантные чанки:")
            if relevant_docs:
                for i, doc in enumerate(relevant_docs):
                    source = doc.metadata.get('source', 'Unknown')
                    # В нашем случае у нас нет номера страницы, только имя файла
                    print(f"Чанк {i+1} (Source: {source}): {doc.page_content[:300]}...") # Печатаем первые 300 символов
            else:
                print("Релевантные чанки не найдены для этого запроса.")
        except Exception as e:
            print(f"[ERROR] Ошибка при выполнении тестового поиска: {e}")

    else: # Если update_knowledge_base вернула False (ошибка)
        print("\n--- Произошла ошибка при обновлении базы знаний. ---")
        # Если произошла ошибка, попробуем хотя бы загрузить старый индекс для поиска
        print("[DEBUG] Попытка загрузить существующий индекс после ошибки обновления.")
        vector_store_object = load_vector_store(INDEX_PERSIST_DIR, embeddings_model)
        if vector_store_object:
            print("[DEBUG] Старый индекс успешно загружен. Можно выполнять поиск по старой базе.")
            # Здесь можно добавить код для тестового поиска по старой базе, если нужно
            # (аналогично блоку выше, но использовать vector_store_object после этой загрузки)
            print("\n--- Тестовый поиск по старой базе ---")
            query = "Что такое RAG?" # Ваш тестовый запрос
            print(f"Запрос: {query}")
            try:
                 relevant_docs = vector_store_object.similarity_search(query, k=3) # Ищем 3 самых похожих чанка
                 print("Найденные релевантные чанки:")
                 if relevant_docs:
                     for i, doc in enumerate(relevant_docs):
                         source = doc.metadata.get('source', 'Unknown')
                         print(f"Чанк {i+1} (Source: {source}): {doc.page_content[:300]}...") # Печатаем первые 300 символов
                 else:
                     print("Релевантные чанки не найдены для этого запроса.")
            except Exception as e:
                 print(f"[ERROR] Ошибка при выполнении тестового поиска по старой базе: {e}")
        else:
            print("[DEBUG] Старый индекс также не найден или не может быть загружен. Поиск невозможен.")
=======
DEEPSEEK_API_KEY = "000"
DEEPSEEK_API_URL = "https://api.deepseek.com/models"
DEEPSEEK_MODEL_NAME = 'deepseek-chat'
EMBEDDING_MODEL_NAME = 'ВАША_МОДЕЛЬ_ЭМБЕДДИНГОВ'
DOCUMENTS_DIR = './documents'
# Путь к папке, где будет сохраняться индекс базы знаний (векторы документов). 📊
# Это чтобы бот "помнил" обработанные документы между запусками.
INDEX_PERSIST_DIR = './vector_index'
ADMIN_USER_IDS = [ 898852116, ]


os.makedirs(DOCUMENTS_DIR, exist_ok=exist_ok) # Создание папки для документов.
os.makedirs(INDEX_PERSIST_DIR, exist_ok=exist_ok) # Создание папки для индекса базы знаний.
>>>>>>> fa6a4e064086aa7f0b7f702b5cdbc1d2b29301a8
